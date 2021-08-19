import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import shared
from docx.shared import Inches

import random
import csv
from calendar import monthrange
import numpy as np


def get_records(
    student_name = "alice",
    month = 6,
    sign_days = 20,
    year=2021
):
    
    workings = [
        ["system test", 1],
        ["measure sample", 1],
        ["calibrate sample", 1],
        ["test device", 1],
        ["data analysis", 0.5],
        ["measure single-qbit gate",1],
        ["measure two-qbit gate",1],
        ["signal calibration",1],
        ["optimize parameters",1],
    ]
    
    
    max_day = monthrange(year, month)[1]
    _dates_pool = np.arange(1,max_day+1,1)

    work_contents_pool = [x[0] for x in workings]
    work_prob = np.asarray([x[1] for x in workings])
    work_prob = work_prob/np.sum(work_prob)

    dates = np.random.choice(
        _dates_pool, size=sign_days, replace=False)
    work_contents = np.random.choice(
        work_contents_pool,size=len(dates),p=work_prob)

    def get_row(
        i, dates, work_contents,
        student_name, month):
        idx_sort = np.argsort(dates)
        idx = idx_sort[i]
        date = dates[idx]
        _content = work_contents[idx]
        if month < 10:
            month = "0"+str(month)
        else:
            month = str(month)
            
        if date < 10:
            date = "0"+str(date)
        else:
            date = str(date)

        return [
            month+"."+date, 
            _content, student_name]

    return [
    get_row(
        i, dates, work_contents,
        student_name, month
        )
    for i in range(len(dates))
]


def set_font(
    paragraph, size, font_name = 'Arial',
    bold = False):
    size = docx.shared.Pt(size)
    for run in paragraph.runs:
        font = run.font
        font.size= size
        font.bold = bold
        font.name = font_name
        
def set_page(document):
    sections = document.sections
    Cm = docx.shared.Cm
    sections[0].page_height = Cm(29.7)
    sections[0].page_width = Cm(21.0)
    for section in sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(0.52)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    

def generate_docx_page(document,records):
    para1 = document.add_paragraph('Sigature form')
    para1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_font(
        para1, 18, font_name = 'Arial',
        bold = True)

    set_page(document)

    n_col = 4
    n_row = 31

    table = document.add_table(rows=1, cols=n_col, style='Table Grid')
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_font = table.style.font
    table_font.name = 'Arial'
    table_font.size = docx.shared.Pt(14)

    table_heads = ['Date','Contents','Name','Sign']
    for i in range(n_col):
        hdr_cells = table.rows[0].cells
        hdr_cells[i].text = table_heads[i]
        for _cell in hdr_cells:
            _cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
            set_font(
                _cell.paragraphs[0], 14, font_name = 'Arial',
                bold = True)

    for _ in range(31):
        table.add_row()

    row_widths = [1.75,7.25,3.25,4.72]
    for i_row,record in enumerate(records):
        row_cells = table.rows[1+i_row].cells
        for i,_text in enumerate(record):
            _cell = row_cells[i]
            _cell.text = _text
            _cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
            _cell.width_rule = docx.enum.table.WD_ROW_HEIGHT_RULE.EXACTLY
            _cell.width = docx.shared.Cm(row_widths[i])

    # set row height
    for i,row in enumerate(table.rows):
        row.height_rule = docx.enum.table.WD_ROW_HEIGHT_RULE.EXACTLY
        row.height = docx.shared.Cm(0.8)
    


if __name__ == '__main__':
    month = 6
    student_names=[
        "Alice",
        "Bob",
        "Martinis",
        "Wallraff",
        "Chow"
    ]
    year = 2021
    sign_days_list = [30,30,27,27,20]

    document = docx.Document()

    for i in range(len(student_names)):
        student_name=student_names[i]
        sign_days = sign_days_list[i]
        records = get_records(
            student_name,month,sign_days,
            year)
        generate_docx_page(document,records)
        if i < len(student_names)-1:
            document.add_page_break()
    document.save(f'{year}-{month}_Auto_sign.docx')
